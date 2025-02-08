import warnings

import yaml
import time
import datetime
import os
import matplotlib.pyplot as plt
import seaborn as sns
import sklearn as sk
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
import logging

pdfmetrics.registerFont(TTFont('TimesNewRoman', '../fonts/Times_New_Roman.ttf'))


class SklearnReportGenerator(sk.base.BaseEstimator):


    def __init__(self, config_file, output_format = "HTML"):
        self.timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        self.report_dir = "reports"
        self.testDir = f"{self.report_dir}/test_{self.timestamp}"
        os.makedirs(self.testDir, exist_ok=True)
        self.config_file = config_file
        self.format = output_format # default - HTML
        self.model = None
        self.metrics = []
        self.pipeline = None
        self._load_config()
        self.selectionParams = False
        self.setLogger()

    def setLogger(self):
        log_file = f"{self.testDir}/files.log"
        log_dir = os.path.dirname(log_file)
        if not os.path.exists(log_dir) and log_dir != "":
            os.makedirs(log_dir)
        logging.basicConfig(
            level=logging.DEBUG,
            format='%(asctime)s - %(levelname)s - %(message)s',
            handlers=[
                logging.FileHandler(log_file),
            ]
        )
        sklearn_logger = logging.getLogger('sklearn')
        sklearn_logger.setLevel(logging.DEBUG)
        def log_warnings(message, category, filename, lineno, file=None, line=None):
            logging.warning(f"{category.__name__}: {message} (in {filename} at line {lineno})")
        warnings.showwarning = log_warnings

    def _load_config(self):
        with open(self.config_file, 'r') as file:
            config = yaml.safe_load(file)

        if config['selectionParams']['enable']:
            self.selectionParams = True

        steps = []
        transformers = []
        for transformer in config['transformers']:
            transformer_class = getattr(sk.preprocessing, transformer['name'])
            transformers.append((transformer['name'], transformer_class(**transformer['params'])))

        for i in transformers:
            steps.append(i)

        self.metrics = config['metrics']
        model_class = self.__get_model_by_name(config['model']['name'])
        self.model = model_class(**config['model']['params'])
        steps.append((f"{self.model.__class__.__name__}", self.model))
        self.pipeline = sk.pipeline.Pipeline(steps)

        if self.selectionParams:
            selectonModel = self.__get_model_by_name(config['selectionParams']['name'])
            self.pipeline = selectonModel(self.pipeline, param_grid = config['selectionParams']['param_grid'], **config['selectionParams']['params'])


    def __get_model_by_name(self, model_name):
        for module_name in dir(sk):
            module = getattr(sk, module_name)
            if isinstance(module, type(sk)):
                for class_name in dir(module):
                    model_class = getattr(module, class_name)
                    if isinstance(model_class, type):
                        if model_name.lower() in class_name.lower():
                            return model_class
        return None

    def fit(self, X, y, X_test=None, y_test=None, *args, **kwargs):
        start_time = time.time()
        self.pipeline.fit(X, y, *args, **kwargs)
        end_time = time.time()
        self.train_time = end_time - start_time

        self.X_train, self.y_train = X, y
        self.X_test, self.y_test = X_test, y_test
        return self

    def predict(self, X):
        y_pred = self.pipeline.predict(X)
        self._generate_report(X, y_pred)
        return y_pred

    def _generate_report(self, X, y_pred):
        y_true = self.y_test
        if sk.base.is_classifier(self.model):
            task_type = "classification"

        report_file = f"{self.testDir}/report_{self.model.__class__.__name__}_{self.timestamp}.html"

        metrics = {}
        for metric in self.metrics:
            metric_func = getattr(sk.metrics, metric['name'])
            params = metric['params']
            metrics[metric['name']] = metric_func(self.y_test, y_pred, **params)

        plots = []
        if "confusion_matrix" in [metric['name'] for metric in self.metrics] and task_type == "classification":
            cm = sk.metrics.confusion_matrix(y_true, y_pred)
            plt.figure(figsize=(6, 4))
            sns.heatmap(cm, annot=True, fmt="d", cmap="Blues")
            plt.xlabel("Predicted")
            plt.ylabel("Actual")
            plt.title("Confusion Matrix")
            plt.savefig(f"{self.report_dir}/confusion_matrix.png")
            plt.close()
            plots.append("confusion_matrix.png")

        if "roc_curve" in [metric['name'] for metric in self.metrics] and task_type == "classification" and len(
                set(y_true)) == 2:
            fpr, tpr, _ = sk.metrics.roc_curve(y_true, y_pred)
            roc_auc = sk.metrics.auc(fpr, tpr)
            plt.figure(figsize=(6, 4))
            plt.plot(fpr, tpr, label=f"AUC = {roc_auc:.2f}")
            plt.plot([0, 1], [0, 1], "r--")
            plt.xlabel("False Positive Rate")
            plt.ylabel("True Positive Rate")
            plt.title("ROC Curve")
            plt.legend()
            plt.savefig(f"{self.report_dir}/roc_curve.png")
            plt.close()
            plots.append("roc_curve.png")

        if self.format == "HTML":
            self.__generate_html_report(report_file, metrics, plots)
        elif self.format == "PDF":
            self.__generate_pdf_report(report_file, metrics, plots)
        elif self.format == "DOCS":
            self.__generate_docx_report(report_file, metrics, plots)


    def __generate_html_report(self, report_file, metrics, plots):
        with open(report_file, "w", encoding="utf-8") as f:
            f.write("<html><head><title>Отчёт о модели</title></head><body>")
            f.write(f"<h1>Отчёт о модели {self.model.__class__.__name__}</h1>")
            f.write(f"<p><b>Время обучения:</b> {self.train_time:.4f} секунд</p>")
            f.write("<h2>Метрики:</h2><ul>")
            for key, value in metrics.items():
                f.write(f"<li><b>{key}:</b> {value:.4f}</li>")
            f.write("</ul>")
            if len(plots) != 0:
                f.write("<h2>Графики:</h2>")
                for plot in plots:
                    f.write(f'<img src="{plot}" width="400px">')
            f.write("</body></html>")

    def __generate_pdf_report(self, report_file, metrics, plots):
        pdf_file = report_file.replace(".html", ".pdf")
        c = canvas.Canvas(pdf_file, pagesize=letter)

        c.setFont("TimesNewRoman", 12)
        c.drawString(100, 750, f"Отчёт о модели {self.model.__class__.__name__}")
        c.drawString(100, 730, f"Время обучения: {self.train_time:.4f} секунд")
        c.drawString(100, 710, "Метрики:")
        y_position = 690
        for key, value in metrics.items():
            c.drawString(100, y_position, f"{key}: {value:.4f}")
            y_position -= 20

        if len(plots) != 0:
            c.drawString(100, y_position, "Графики:")
            y_position -= 20
            for plot in plots:
                c.drawString(100, y_position, f"{plot}")
                y_position -= 20
        c.save()

    def __generate_docx_report(self, report_file, metrics, plots):
        docx_file = report_file.replace(".html", ".docx")
        doc = Document()

        doc.add_heading(f"Отчёт о модели {self.model.__class__.__name__}", 0)
        doc.add_paragraph(f"Время обучения: {self.train_time:.4f} секунд")

        doc.add_heading("Метрики:", level=1)
        for key, value in metrics.items():
            doc.add_paragraph(f"{key}: {value:.4f}")

        if len(plots) != 0:
            doc.add_heading("Графики:", level=1)
            for plot in plots:
                doc.add_paragraph(f"{plot}")

        doc.save(docx_file)
