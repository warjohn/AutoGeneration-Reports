from docx import Document


class GenerateDocsReport():

    @classmethod
    def generateReport(cls, report_file, model_name, train_time, result_metrics):
        docx_file = report_file.replace(".html", ".docx")
        doc = Document()

        doc.add_heading(f"Отчёт о модели {model_name}", 0)
        doc.add_paragraph(f"Время обучения: {train_time:.4f} секунд")

        doc.add_heading("Метрики:", level=1)
        for key, value in result_metrics.items():
            doc.add_paragraph(f"{key}: {value:.4f}")

        doc.save(docx_file)
